from __future__ import annotations

import json
import os
import queue
import socket
import sys
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from html import escape
import zipfile
import xml.etree.ElementTree as ET
import tempfile
# Try to use lxml for more robust XML parsing if available
try:
    import lxml.etree as LET  # type: ignore
except Exception:
    LET = None  # type: ignore
try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:
    DocxDocument = None  # type: ignore

from PySide6.QtCore import QObject, Qt, QTimer, Signal, QEvent, QMarginsF
from PySide6.QtGui import QAction, QCloseEvent, QActionGroup, QKeySequence, QTextListFormat, QTextCharFormat, QColor, QPageLayout, QPageSize, QTextOption
from PySide6.QtWidgets import (
    QApplication, QFileDialog, QMainWindow, QTextEdit, QMessageBox,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QDockWidget, QFrame
)
from PySide6.QtPrintSupport import QPrinter

# Ensure we can import gui_core from the Walls directory

def find_walls_root(start: Optional[Path] = None) -> Optional[Path]:
    p = (start or Path(__file__).resolve()).parent
    for _ in range(8):
        if p.name == "Walls" and p.is_dir():
            return p
        p = p.parent
    return None

WALLS_ROOT = find_walls_root(Path(__file__).resolve())
if WALLS_ROOT and str(WALLS_ROOT) not in sys.path:
    sys.path.insert(0, str(WALLS_ROOT))

# Optional: apply theme if available
try:
    from gui_core.apply_theme import apply_theme as apply_app_theme  # type: ignore
    from gui_core.components.toolbar.widgets import ToolBar
    from gui_core.components.breadcrumbs.widgets import Breadcrumbs
    from gui_core.components.line_edit.widgets import LineEdit
    from gui_core.components.cards.widgets import Card
except Exception:
    def apply_app_theme(_app):
        pass
    # fallbacks if components missing
    from PySide6.QtWidgets import QToolBar as ToolBar, QLineEdit as LineEdit
    class Card(QFrame):
        def __init__(self, title: str = "", subtitle: str = "", parent=None):
            super().__init__(parent)
            self.setObjectName("cardContainer")
            self._lay = QVBoxLayout(self)
            if title:
                self._lay.addWidget(QLabel(title))
            if subtitle:
                lb = QLabel(subtitle); lb.setWordWrap(True); self._lay.addWidget(lb)
        def addWidget(self, w):
            self._lay.addWidget(w)
        def addActionButton(self, btn):
            self._lay.addWidget(btn)
    class Breadcrumbs(QWidget):
        def __init__(self, parts=None, parent=None):
            super().__init__(parent)
            lay = QHBoxLayout(self); lay.setContentsMargins(0,0,0,0)
            self._label = QLabel(" / ".join(parts or ["Home"]))
            lay.addWidget(self._label)

# Import icons helper
from gui.custom_widgets import icon

# Import Rust core (PyO3) with a Python fallback for immediate usability
try:
    import word_core  # type: ignore
except Exception:
    word_core = None


class PyFallbackDocument:
    def __init__(self) -> None:
        self._text = ""

    def set_text(self, text: str) -> None:
        self._text = text

    def insert_text(self, offset: int, text: str) -> None:
        if offset < 0:
            offset = 0
        if offset > len(self._text):
            offset = len(self._text)
        self._text = self._text[:offset] + text + self._text[offset:]

    def get_text(self) -> str:
        return self._text

    def clear(self) -> None:
        self._text = ""

    def open(self, path: str) -> None:
        with open(path, "r", encoding="utf-8") as f:
            self._text = f.read()

    def save(self, path: str) -> None:
        with open(path, "w", encoding="utf-8") as f:
            f.write(self._text)


DocumentClass = getattr(word_core, "Document", PyFallbackDocument) if word_core else PyFallbackDocument


@dataclass
class Command:
    name: str
    args: dict


class CommandBridge(QObject):
    command_received = Signal(object)


class CommandServer(threading.Thread):
    def __init__(self, host: str, port: int, bridge: CommandBridge, port_file: Optional[Path] = None):
        super().__init__(daemon=True)
        self._host = host
        self._port = port
        self._bridge = bridge
        self._sock: Optional[socket.socket] = None
        self._stop_event = threading.Event()
        self._port_file = Path(port_file) if port_file else None

    def stop(self):
        self._stop_event.set()
        try:
            if self._sock:
                self._sock.close()
        except Exception:
            pass

    def _write_port_file(self):
        if not self._port_file:
            return
        try:
            self._port_file.write_text(str(self._port), encoding='utf-8')
        except Exception:
            pass

    def run(self):
        try:
            self._sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            self._sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            # Try to bind to requested port; if in use, try the next few ports
            bound = False
            for p in range(self._port, self._port + 20):
                try:
                    self._sock.bind((self._host, p))
                    self._port = p
                    bound = True
                    break
                except OSError:
                    continue
            if not bound:
                print("CLI server failed to start: no available port in range")
                return
            self._sock.listen(1)
            print(f"CLI server listening on {self._host}:{self._port}")
            self._write_port_file()
        except Exception as e:
            print(f"CLI server failed to start: {e}")
            return

        while not self._stop_event.is_set():
            try:
                self._sock.settimeout(0.5)
                try:
                    conn, _addr = self._sock.accept()
                except socket.timeout:
                    continue

                with conn:
                    buf = b""
                    while not self._stop_event.is_set():
                        chunk = conn.recv(4096)
                        if not chunk:
                            break
                        buf += chunk
                        while b"\n" in buf:
                            line, buf = buf.split(b"\n", 1)
                            line = line.strip()
                            if not line:
                                continue
                            try:
                                data = json.loads(line.decode("utf-8"))
                                cmd = Command(name=data.get("cmd", ""), args=data.get("args", {}))
                                self._bridge.command_received.emit(cmd)
                                conn.sendall(b"{\"ok\":true}\n")
                            except Exception as e:
                                err = json.dumps({"ok": False, "error": str(e)}) + "\n"
                                conn.sendall(err.encode("utf-8"))
            except Exception:
                continue


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Words - Editor")

        # Document backend
        self.doc = DocumentClass()

        # Central area with header and editor
        central = QWidget(self)
        vbox = QVBoxLayout(central)
        vbox.setContentsMargins(16, 12, 16, 12)
        vbox.setSpacing(10)

        # Header: breadcrumbs on top, then title and subtitle
        header = QWidget(central)
        hv = QVBoxLayout(header)
        hv.setContentsMargins(0, 0, 0, 0)
        hv.setSpacing(6)
        # Replace book-specific header with a persistent formatting toolbar
        self.header_tb = ToolBar("Format", header)
        hv.addWidget(self.header_tb)
        vbox.addWidget(header)

        # Editor - simplified classical layout
        self.editor = QTextEdit(central)
        self.editor.setObjectName("wordsEditor")
        # Ensure rich text is accepted for HTML rendering
        self.editor.setAcceptRichText(True)
        # Make text wrap to the editor width and add comfortable margins
        self.editor.setLineWrapMode(QTextEdit.WidgetWidth)
        self.editor.setWordWrapMode(QTextOption.WordWrap)
        self.editor.document().setDocumentMargin(24)
        # Default CSS for consistent layout/typography when rendering HTML content
        self.default_css = (
            "body{font-family:-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;font-size:14pt;line-height:1.6;margin:0;padding:0;}"
            "p{margin:0 0 12px;} h1{font-size:1.8em;margin:0 0 16px;} h2{font-size:1.6em;margin:0 0 14px;} "
            "h3{font-size:1.4em;margin:0 0 12px;} h4{font-size:1.2em;margin:0 0 10px;} h5{font-size:1.1em;margin:0 0 8px;} h6{font-size:1.0em;margin:0 0 6px;} "
            "ul,ol{margin:0 0 12px 24px;} li{margin:4px 0;} a{color:#0a84ff;text-decoration:none;} a:hover{text-decoration:underline;}"
            "table{border-collapse:collapse;width:100%;margin:0 0 12px;} th,td{border:1px solid #ddd;padding:6px;vertical-align:top;}"
        )
        # Track whether the editor currently displays rich HTML content
        self.is_rich_mode = False
        vbox.addWidget(self.editor, 1)
        central.setLayout(vbox)
        self.setCentralWidget(central)
        self.editor.textChanged.connect(self._on_editor_changed)

        # Top file toolbar (keep minimal per theme)
        self.file_tb = ToolBar("File", self)
        self.addToolBar(Qt.TopToolBarArea, self.file_tb)
        # Match format toolbar icon size with file toolbar
        try:
            self.header_tb.setIconSize(self.file_tb.iconSize())
        except Exception:
            pass
        act_open = self.file_tb.addAction(icon("folder.svg"), "")
        act_save = self.file_tb.addAction(icon("floppy-disk.svg"), "")
        act_open.triggered.connect(self.open_file_dialog)
        act_save.triggered.connect(self.save_file_dialog)

        # Floating formatting bubble inspired by CMS editor
        self.format_bubble = Card(title="", subtitle="")
        self.format_bubble.setObjectName("formatBubble")
        # Card already owns a layout in gui_core; just add the toolbar as its content
        self.bubble_tb = ToolBar("inline", self)
        self.bubble_tb.setFloatable(False); self.bubble_tb.setMovable(False)
        self.bubble_tb.setIconSize(self.file_tb.iconSize())
        self.bubble_tb.setStyleSheet("")
        # Actions on bubble
        self.act_bold = self.bubble_tb.addAction(icon("bold.svg"), "")
        self.act_italic = self.bubble_tb.addAction(icon("italic.svg"), "")
        self.act_underline = self.bubble_tb.addAction(icon("underline.svg"), "")
        self.bubble_tb.addSeparator()
        self.act_color = self.bubble_tb.addAction(icon("colors.svg"), "")
        self.bubble_tb.addSeparator()
        self.act_align_left = self.bubble_tb.addAction(icon("align-left.svg"), "")
        self.act_align_center = self.bubble_tb.addAction(icon("align-center.svg"), "")
        self.act_align_right = self.bubble_tb.addAction(icon("align-right.svg"), "")
        self.act_align_justify = self.bubble_tb.addAction(icon("align-justify.svg"), "")
        self.bubble_tb.addSeparator()
        self.act_list_ul = self.bubble_tb.addAction(icon("list-ul.svg"), "")
        self.act_list_ol = self.bubble_tb.addAction(icon("list-ol.svg"), "")
        self.bubble_tb.addSeparator()
        self.act_link = self.bubble_tb.addAction(icon("link.svg"), "")
        self.act_image = self.bubble_tb.addAction(icon("image.svg"), "")
        self.act_code = self.bubble_tb.addAction(icon("code.svg"), "")
        # Mirror actions on persistent header toolbar
        if hasattr(self, 'header_tb'):
            ht = self.header_tb
            ht.clear()
            ht.addAction(self.act_bold)
            ht.addAction(self.act_italic)
            ht.addAction(self.act_underline)
            ht.addSeparator()
            ht.addAction(self.act_color)
            ht.addSeparator()
            ht.addAction(self.act_align_left)
            ht.addAction(self.act_align_center)
            ht.addAction(self.act_align_right)
            ht.addAction(self.act_align_justify)
            ht.addSeparator()
            ht.addAction(self.act_list_ul)
            ht.addAction(self.act_list_ol)
            ht.addSeparator()
            ht.addAction(self.act_link)
            ht.addAction(self.act_image)
            ht.addAction(self.act_code)
        # Make format actions reflect state
        self.act_bold.setCheckable(True)
        self.act_italic.setCheckable(True)
        self.act_underline.setCheckable(True)
        self.act_code.setCheckable(True)
        # Keyboard shortcuts available globally
        try:
            self.act_bold.setShortcut(QKeySequence.Bold)
            self.act_italic.setShortcut(QKeySequence.Italic)
            self.act_underline.setShortcut(QKeySequence.Underline)
        except Exception:
            pass
        for a in (self.act_bold, self.act_italic, self.act_underline):
            self.addAction(a)
        # Group alignment actions exclusively
        self.align_group = QActionGroup(self)
        for a in (self.act_align_left, self.act_align_center, self.act_align_right, self.act_align_justify):
            a.setCheckable(True)
            self.align_group.addAction(a)
        self.align_group.setExclusive(True)
        # Helpful tooltips
        accel_b = "Cmd+B" if sys.platform == "darwin" else "Ctrl+B"
        accel_i = "Cmd+I" if sys.platform == "darwin" else "Ctrl+I"
        accel_u = "Cmd+U" if sys.platform == "darwin" else "Ctrl+U"
        self.act_bold.setToolTip(f"Bold ({accel_b})")
        self.act_italic.setToolTip(f"Italic ({accel_i})")
        self.act_underline.setToolTip(f"Underline ({accel_u})")
        self.act_color.setToolTip("Text color")
        self.act_code.setToolTip("Inline code")
        self.act_align_left.setToolTip("Align left")
        self.act_align_center.setToolTip("Align center")
        self.act_align_right.setToolTip("Align right")
        self.act_align_justify.setToolTip("Justify")
        self.act_list_ul.setToolTip("Bullet list (simple)")
        self.act_list_ol.setToolTip("Numbered list (simple)")
        self.act_link.setToolTip("Insert link")
        self.act_image.setToolTip("Insert image")
        try:
            self.format_bubble.addWidget(self.bubble_tb)
        except Exception:
            # Fallback for local Card stub
            lay = getattr(self.format_bubble, 'layout', None)
            if callable(lay):
                self.format_bubble.layout().addWidget(self.bubble_tb)
        # Keep bubble permanently hidden and detached
        try:
            self.format_bubble.setVisible(False)
            self.format_bubble.hide()
            self.format_bubble.setParent(None)
        except Exception:
            pass

        # Connect actions from bubble
        self.act_bold.triggered.connect(self._toggle_bold)
        self.act_italic.triggered.connect(self._toggle_italic)
        self.act_underline.triggered.connect(self._toggle_underline)
        self.act_color.triggered.connect(self._pick_color)
        self.act_align_left.triggered.connect(lambda: self.editor.setAlignment(Qt.AlignLeft))
        self.act_align_center.triggered.connect(lambda: self.editor.setAlignment(Qt.AlignHCenter))
        self.act_align_right.triggered.connect(lambda: self.editor.setAlignment(Qt.AlignRight))
        self.act_align_justify.triggered.connect(lambda: self.editor.setAlignment(Qt.AlignJustify))
        self.act_list_ul.triggered.connect(self._make_bullet_list)
        self.act_list_ol.triggered.connect(self._make_number_list)
        self.act_link.triggered.connect(self._insert_link)
        self.act_image.triggered.connect(self._insert_image)
        self.act_code.triggered.connect(self._toggle_code)

        # Disable floating bubble on selection changes (user request)
        # We intentionally DO NOT connect selection/cursor signals and DO NOT install
        # any event filter that would show or position the bubble.
        # (No-op here prevents runtime warnings from disconnect attempts.)

        # Document panel removed - using classical text editor approach


        # CLI bridge using shared server
        self.bridge = CommandBridge()
        self.bridge.command_received.connect(self.handle_command)
        
        # Register with shared server instead of creating own server
        try:
            from shared_server import get_shared_server
            shared_server = get_shared_server()
            
            # Create command handler function for shared server
            def command_handler(cmd: str, args: dict) -> dict:
                print(f"Words GUI received command: {cmd} with args: {args}")
                try:
                    print("command_handler: creating command object")
                    command = Command(name=cmd, args=args)
                    print("command_handler: emitting signal to GUI thread")
                    # Emit directly; cross-thread signal will be queued to the GUI thread
                    self.bridge.command_received.emit(command)
                    print("command_handler: signal emitted, returning success")
                    return {'status': 'success', 'message': f'Command {cmd} executed'}
                except Exception as e:
                    print(f"Error in command_handler: {e}")
                    return {'status': 'error', 'message': str(e)}
            
            self.assigned_port = shared_server.register_app("words", command_handler)
            
            # Write port file for CLI compatibility
            port_file = Path(tempfile.gettempdir()) / "words_gui_port"
            with open(port_file, 'w') as f:
                f.write(str(self.assigned_port))
            
            print(f"Port file written to: {port_file}")
            
            print(f"Words registered with shared server on port {self.assigned_port}")
        except Exception as e:
            print(f"Failed to register with shared server: {e}")
            # Fallback to original CommandServer
            port_hint = 8765
            port_file = Path(tempfile.gettempdir()) / "words_gui_port"
            self.server = CommandServer("127.0.0.1", port_hint, self.bridge, port_file=port_file)
            self.server.start()
            self.assigned_port = None

        # Try to apply theme
        apply_app_theme(QApplication.instance())

    def _on_editor_changed(self):
        # Push edits to backend
        text = self.editor.toPlainText()
        try:
            self.doc.set_text(text)
        except Exception as e:
            print("Backend set_text error:", e)

    def set_editor_text(self, text: str):
        # Avoid recursive signals by blocking
        self.editor.blockSignals(True)
        # If the input looks like HTML or we're in rich mode, set as HTML
        t = text or ""
        looks_like_html = ("<html" in t.lower()) or ("</" in t) or ("<p" in t.lower()) or ("<h1" in t.lower()) or ("<b>" in t.lower()) or ("<i>" in t.lower()) or ("<u>" in t.lower()) or ("<table" in t.lower())
        if looks_like_html:
            # Auto-enable rich mode when HTML content is detected
            self.is_rich_mode = True
            # Sanitize HTML for QTextEdit compatibility
            sanitized_html = self._sanitize_html_for_qtextedit(t)
            self.editor.setHtml(sanitized_html)
        elif getattr(self, 'is_rich_mode', False):
            sanitized_html = self._sanitize_html_for_qtextedit(t)
            self.editor.setHtml(sanitized_html)
        else:
            self.editor.setPlainText(t)
        self.editor.blockSignals(False)

    def _sanitize_html_for_qtextedit(self, html: str) -> str:
        """Sanitize HTML content to be compatible with QTextEdit's limited HTML subset."""
        if not html:
            return html
        
        # Remove DOCTYPE declarations and full HTML document structure
        # QTextEdit works better with HTML fragments
        import re
        
        # Remove DOCTYPE declaration
        html = re.sub(r'<!DOCTYPE[^>]*>', '', html, flags=re.IGNORECASE)
        
        # Extract body content if it's a full HTML document
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL | re.IGNORECASE)
        if body_match:
            html = body_match.group(1)
        
        # Remove html and head tags if present
        html = re.sub(r'</?html[^>]*>', '', html, flags=re.IGNORECASE)
        html = re.sub(r'<head>.*?</head>', '', html, flags=re.DOTALL | re.IGNORECASE)
        
        # Clean up whitespace
        html = html.strip()
        
        # If we have content, wrap it with basic styling
        if html:
            # Add basic CSS styling inline since QTextEdit has limited CSS support
            styled_html = f'<div style="font-family: {self.editor.font().family()}; font-size: {self.editor.font().pointSize()}pt;">{html}</div>'
            return styled_html
        
        return html

    # --- Formatting helpers and stats ---
    def _toggle_bold(self):
        fmt = self.editor.currentCharFormat()
        is_bold = fmt.fontWeight() > 50
        fmt.setFontWeight(400 if is_bold else 700)
        self.editor.mergeCurrentCharFormat(fmt)

    def _toggle_italic(self):
        fmt = self.editor.currentCharFormat()
        fmt.setFontItalic(not fmt.fontItalic())
        self.editor.mergeCurrentCharFormat(fmt)

    def _toggle_underline(self):
        fmt = self.editor.currentCharFormat()
        fmt.setFontUnderline(not fmt.fontUnderline())
        self.editor.mergeCurrentCharFormat(fmt)

    def _make_bullet_list(self):
        cursor = self.editor.textCursor()
        if cursor.currentList():
            # toggle off list
            lst = cursor.currentList()
            block = cursor.block()
            lst.remove(block)
        else:
            fmt = QTextListFormat()
            fmt.setStyle(QTextListFormat.ListDisc)
            cursor.createList(fmt)

    def _make_number_list(self):
        cursor = self.editor.textCursor()
        if cursor.currentList():
            lst = cursor.currentList()
            block = cursor.block()
            lst.remove(block)
        else:
            fmt = QTextListFormat()
            fmt.setStyle(QTextListFormat.ListDecimal)
            cursor.createList(fmt)

    def _pick_color(self):
        from PySide6.QtWidgets import QColorDialog
        color = QColorDialog.getColor(QColor("#212121"), self, "Select text color")
        if not color.isValid():
            return
        fmt = QTextCharFormat()
        fmt.setForeground(color)
        self.editor.mergeCurrentCharFormat(fmt)

    def _insert_link(self):
        from PySide6.QtWidgets import QInputDialog
        cursor = self.editor.textCursor()
        if not cursor.hasSelection():
            return
        url, ok = QInputDialog.getText(self, "Insert link", "URL:")
        if not ok or not url:
            return
        fmt = QTextCharFormat()
        fmt.setAnchor(True)
        fmt.setAnchorHref(url)
        fmt.setForeground(QColor("#1F2937"))
        fmt.setFontUnderline(True)
        self.editor.mergeCurrentCharFormat(fmt)

    def _insert_image(self):
        path, _ = QFileDialog.getOpenFileName(self, "Insert image", os.getcwd(), "Images (*.png *.jpg *.jpeg *.gif *.webp *.bmp)")
        if not path:
            return
        from PySide6.QtGui import QTextImageFormat
        img = QTextImageFormat()
        img.setName(path)
        cursor = self.editor.textCursor()
        cursor.insertImage(img)

    def _toggle_code(self):
        fmt = self.editor.currentCharFormat()
        is_mono = fmt.fontFamily().lower().startswith("menlo") or fmt.fontFixedPitch()
        if is_mono:
            fmt.setFontFamily("")
            fmt.setFontFixedPitch(False)
            fmt.clearBackground()
        else:
            fmt.setFontFamily("Menlo")
            fmt.setFontFixedPitch(True)
            fmt.setBackground(QColor("#F3F4F6"))
        self.editor.mergeCurrentCharFormat(fmt)

    def _print_document(self, printer: QPrinter):
        try:
            doc = self.editor.document()
            doc.print(printer)
        except Exception:
            pass

    def _update_preview(self):
        try:
            if getattr(self, 'preview', None) and getattr(self, 'preview_dock', None) and self.preview_dock.isVisible():
                self.preview.updatePreview()
        except Exception:
            pass

    def open_file_dialog(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open Document", os.getcwd(), "Documents (*.txt *.docx *.odt)")
        if not path:
            return
        self.open_path(path)

    def open_path(self, path: str):
        ext = Path(path).suffix.lower().lstrip('.')
        if ext == 'docx':
            # Prefer python-docx for better fidelity (headings, bold/italic/underline, tables)
            if DocxDocument:
                try:
                    from docx.oxml.text.paragraph import CT_P  # type: ignore
                    from docx.oxml.table import CT_Tbl  # type: ignore
                    from docx.text.paragraph import Paragraph as DocxParagraph  # type: ignore
                    from docx.table import Table as DocxTable  # type: ignore
                except Exception:
                    CT_P = None  # type: ignore
                    CT_Tbl = None  # type: ignore
                    DocxParagraph = None  # type: ignore
                    DocxTable = None  # type: ignore
                try:
                    doc = DocxDocument(path)

                    def para_to_html(p) -> str:
                        tag = "p"
                        try:
                            sty = (p.style.name if p.style is not None else "") or ""
                        except Exception:
                            sty = ""
                        if sty.startswith("Heading"):
                            lvl = ''.join(ch for ch in sty if ch.isdigit()) or "1"
                            tag = f"h{lvl}"
                        parts_inner: list[str] = []
                        for r in getattr(p, 'runs', []) or []:
                            t = escape(getattr(r, 'text', '') or "")
                            if not t:
                                continue
                            t = t.replace("\n", "<br/>")
                            if bool(getattr(r, 'bold', False)):
                                t = f"<b>{t}</b>"
                            if bool(getattr(r, 'italic', False)):
                                t = f"<i>{t}</i>"
                            if bool(getattr(r, 'underline', False)):
                                t = f"<u>{t}</u>"
                            parts_inner.append(t)
                        inner = "".join(parts_inner) or "<br/>"
                        return f"<{tag}>" + inner + f"</{tag}>"

                    def table_to_html(t) -> str:
                        rows_html: list[str] = []
                        for row in getattr(t, 'rows', []) or []:
                            cells_html: list[str] = []
                            for cell in getattr(row, 'cells', []) or []:
                                # A cell can contain multiple paragraphs
                                para_htmls = [para_to_html(p) for p in getattr(cell, 'paragraphs', []) or []]
                                cell_html = "<br/>".join(h.replace('<p>', '').replace('</p>', '') for h in para_htmls)
                                cells_html.append(f"<td>{cell_html}</td>")
                            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
                        return "<table>" + "".join(rows_html) + "</table>"

                    html_parts: list[str] = []
                    # Try to preserve body order if internals available
                    if DocxParagraph and DocxTable and hasattr(doc, 'element') and hasattr(doc.element, 'body'):
                        try:
                            for child in doc.element.body.iterchildren():
                                tag = getattr(child, 'tag', '')
                                if CT_Tbl is not None and isinstance(child, CT_Tbl):
                                    html_parts.append(table_to_html(DocxTable(child, doc._body)))
                                elif CT_P is not None and isinstance(child, CT_P):
                                    html_parts.append(para_to_html(DocxParagraph(child, doc._body)))
                                else:
                                    # Fallback: skip unknown block types
                                    pass
                        except Exception:
                            # Fallback to paragraphs then tables if order traversal failed
                            html_parts.extend(para_to_html(p) for p in getattr(doc, 'paragraphs', []) or [])
                            html_parts.extend(table_to_html(t) for t in getattr(doc, 'tables', []) or [])
                    else:
                        # Fallback simple traversal
                        html_parts.extend(para_to_html(p) for p in getattr(doc, 'paragraphs', []) or [])
                        html_parts.extend(table_to_html(t) for t in getattr(doc, 'tables', []) or [])

                    body = "\n".join(html_parts) if html_parts else "<p></p>"
                    # Use sanitized HTML for QTextEdit compatibility
                    self.editor.blockSignals(True)
                    self.set_editor_text(body)
                    self.editor.blockSignals(False)
                    self.is_rich_mode = True
                    print(f"Rendered DOCX via python-docx; HTML length={len(html)}")
                    try:
                        self.doc.set_text(self.editor.toPlainText())
                    except Exception as e:
                        print("Backend set_text error:", e)
                    return
                except Exception as e:
                    print("DOCX render via python-docx failed:", e)
            # If no python-docx or it failed, fall back to zip-first parsing
            # Try robust zip/XML parsing first
            try:
                with zipfile.ZipFile(path, 'r') as z:
                    xml_bytes = z.read('word/document.xml')

                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                if LET is not None:
                    try:
                        parser = LET.XMLParser(recover=True)
                        root = LET.fromstring(xml_bytes, parser=parser)
                        use_lxml = True
                    except Exception:
                        root = ET.fromstring(xml_bytes)
                        use_lxml = False
                else:
                    root = ET.fromstring(xml_bytes)
                    use_lxml = False

                def findall(elem, p: str):
                    return elem.findall(p, ns) if not use_lxml else elem.findall(p, namespaces=ns)
                def find(elem, p: str):
                    return elem.find(p, ns) if not use_lxml else elem.find(p, namespaces=ns)

                parts: list[str] = []
                for p in findall(root, './/w:p'):
                    tag = 'p'
                    pPr = find(p, 'w:pPr')
                    if pPr is not None:
                        pStyle = find(pPr, 'w:pStyle')
                        if pStyle is not None:
                            w_ns = ns['w']
                            val = pStyle.get(f'{{{w_ns}}}val') or ''
                            if isinstance(val, bytes):
                                val = val.decode('utf-8', errors='ignore')
                            if val.startswith('Heading'):
                                lvl = ''.join(ch for ch in val if ch.isdigit()) or '1'
                                tag = f'h{lvl}'
                    buf_parts: list[str] = []
                    for r in findall(p, 'w:r'):
                        rPr = find(r, 'w:rPr')
                        bold = bool(rPr is not None and find(rPr, 'w:b') is not None)
                        italic = bool(rPr is not None and find(rPr, 'w:i') is not None)
                        underline = bool(rPr is not None and find(rPr, 'w:u') is not None)
                        texts = findall(r, 'w:t')
                        if not texts:
                            for _ in findall(r, 'w:br'):
                                buf_parts.append('<br/>')
                            continue
                        chunks = []
                        for tnode in texts:
                            tval = tnode.text if not use_lxml else (tnode.text or '')
                            if isinstance(tval, bytes):
                                tval = tval.decode('utf-8', errors='ignore')
                            chunks.append(tval or '')
                        chunk = ''.join(chunks)
                        ttxt = escape(chunk).replace('\n', '<br/>')
                        if bold:
                            ttxt = f'<b>{ttxt}</b>'
                        if italic:
                            ttxt = f'<i>{ttxt}</i>'
                        if underline:
                            ttxt = f'<u>{ttxt}</u>'
                        buf_parts.append(ttxt)
                    inner = ''.join(buf_parts) if buf_parts else '<br/>'
                    parts.append(f'<{tag}>' + inner + f'</{tag}>')
                body = '\n'.join(parts) if parts else '<p></p>'
                html = '<!DOCTYPE html>\n<html><head><meta charset="utf-8"><style>' + self.default_css + '</style></head><body>' + body + '</body></html>'
                self.editor.blockSignals(True)
                self.editor.setHtml(html)
                self.editor.blockSignals(False)
                self.is_rich_mode = True
                print(f"Rendered DOCX via zip-first; HTML length={len(html)}")
                try:
                    self.doc.set_text(self.editor.toPlainText())
                except Exception as e:
                    print('Backend set_text error:', e)
                return
            except Exception as zip_e:
                print('DOCX zip-first parse failed, trying python-docx:', zip_e)
                # Fall through to python-docx
            if DocxDocument:
                try:
                    doc = DocxDocument(path)
                    parts: list[str] = []
                    for p in doc.paragraphs:
                        tag = "p"
                        try:
                            sty = (p.style.name if p.style is not None else "") or ""
                        except Exception:
                            sty = ""
                        if sty.startswith("Heading"):
                            lvl = ''.join(ch for ch in sty if ch.isdigit()) or "1"
                            tag = f"h{lvl}"
                        buf_parts: list[str] = []
                        for r in p.runs:
                            t = escape(r.text or "")
                            if not t:
                                continue
                            t = t.replace("\n", "<br/>")
                            if getattr(r, 'bold', False):
                                t = f"<b>{t}</b>"
                            if getattr(r, 'italic', False):
                                t = f"<i>{t}</i>"
                            if getattr(r, 'underline', False):
                                t = f"<u>{t}</u>"
                            buf_parts.append(t)
                        inner = "".join(buf_parts) or "<br/>"
                        parts.append(f"<{tag}>" + inner + f"</{tag}>")
                    body = "\n".join(parts) if parts else "<p></p>"
                    html = (
                        "<!DOCTYPE html>\n"
                        "<html><head><meta charset='utf-8'><style>" + self.default_css + "</style></head><body>" + body + "</body></html>"
                    )
                    self.editor.blockSignals(True)
                    self.editor.setHtml(html)
                    self.editor.blockSignals(False)
                    self.is_rich_mode = True
                    print(f"Rendered DOCX via python-docx; HTML length={len(html)}")
                    try:
                        self.doc.set_text(self.editor.toPlainText())
                    except Exception as e:
                        print("Backend set_text error:", e)
                    return
                except Exception as e:
                    print("DOCX render via python-docx failed:", e)
            # Fallback to plain text backend open
            try:
                self.doc.open(path)
                self.is_rich_mode = False
                self.set_editor_text(self.doc.get_text())
            except Exception as e:
                QMessageBox.critical(self, "Open Error", str(e))
            return
        elif ext == 'odt':
            # Try using the enhanced Rust parser first
            if word_core and hasattr(word_core, 'read_odt_structured_json'):
                try:
                    structured_json = word_core.read_odt_structured_json(path)
                    structured_data = json.loads(structured_json)
                    
                    # Convert structured data to HTML
                    html_parts = []
                    for element in structured_data.get('elements', []):
                        # Handle the actual JSON structure from Rust parser
                        if 'Heading' in element:
                            heading_data = element['Heading']
                            level = heading_data.get('level', 1)
                            level = max(1, min(6, level))
                            
                            # Process runs (text segments with styles)
                            formatted_parts = []
                            for run in heading_data.get('runs', []):
                                text = run.get('text', '')
                                style = run.get('style', {})
                                
                                formatted_text = escape(text).replace('\n', '<br/>')
                                if style.get('bold', False):
                                    formatted_text = f"<b>{formatted_text}</b>"
                                if style.get('italic', False):
                                    formatted_text = f"<i>{formatted_text}</i>"
                                if style.get('underline', False):
                                    formatted_text = f"<u>{formatted_text}</u>"
                                formatted_parts.append(formatted_text)
                            
                            content = ''.join(formatted_parts)
                            html_parts.append(f'<h{level}>{content}</h{level}>')
                            
                        elif 'Paragraph' in element:
                            paragraph_data = element['Paragraph']
                            
                            # Process runs (text segments with styles)
                            formatted_parts = []
                            for run in paragraph_data.get('runs', []):
                                text = run.get('text', '')
                                style = run.get('style', {})
                                
                                formatted_text = escape(text).replace('\n', '<br/>')
                                if style.get('bold', False):
                                    formatted_text = f"<b>{formatted_text}</b>"
                                if style.get('italic', False):
                                    formatted_text = f"<i>{formatted_text}</i>"
                                if style.get('underline', False):
                                    formatted_text = f"<u>{formatted_text}</u>"
                                formatted_parts.append(formatted_text)
                            
                            content = ''.join(formatted_parts)
                            html_parts.append(f'<p>{content}</p>')
                            
                        elif 'List' in element:
                            list_data = element['List']
                            is_ordered = list_data.get('ordered', False)
                            list_tag = 'ol' if is_ordered else 'ul'
                            
                            # For now, skip empty lists or lists without proper item structure
                            # The Rust parser seems to have issues with list item extraction
                            # This prevents empty <ul></ul> or <ol></ol> from appearing
                            items = list_data.get('items', [])
                            if items and any(items):  # Only render if there are actual items
                                html_parts.append(f'<{list_tag}></{list_tag}>')
                            # Note: List items are typically parsed as separate Paragraph elements
                            # so they will be handled by the Paragraph case above
                        else:
                            # Handle unknown element types - skip them or try to extract text
                            # Don't display raw JSON, just skip unknown elements
                            continue
                    
                    body_html = '\n'.join(html_parts) if html_parts else '<p></p>'
                    
                    # Set the HTML content
                    self.editor.blockSignals(True)
                    self.set_editor_text(body_html)
                    self.editor.blockSignals(False)
                    self.is_rich_mode = True
                    print(f"Rendered ODT via Rust parser; HTML length={len(body_html)}")
                    
                    try:
                        self.doc.set_text(self.editor.toPlainText())
                    except Exception as e:
                        print('Backend set_text error:', e)
                    return
                except Exception as e:
                    print('Enhanced ODT parsing failed, falling back to legacy parser:', e)
            
            # Fallback to legacy ODT parsing
            try:
                with zipfile.ZipFile(path, 'r') as z:
                    xml_bytes = z.read('content.xml')
                if LET is not None:
                    try:
                        parser = LET.XMLParser(recover=True)
                        root = LET.fromstring(xml_bytes, parser=parser)
                        use_lxml = True
                    except Exception:
                        root = ET.fromstring(xml_bytes)
                        use_lxml = False
                else:
                    root = ET.fromstring(xml_bytes)
                    use_lxml = False
                ns = {
                    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
                }
                def find(elem, p: str):
                    return elem.find(p, ns) if not use_lxml else elem.find(p, namespaces=ns)
                def findall(elem, p: str):
                    return elem.findall(p, ns) if not use_lxml else elem.findall(p, namespaces=ns)
                # Navigate to office:text
                body = find(root, './/office:body')
                office_text = find(body, 'office:text') if body is not None else None
                if office_text is None:
                    raise ValueError('Invalid ODT: missing office:text')

                def extract_text(node) -> str:
                    parts: list[str] = []
                    try:
                        if node.text:
                            parts.append(node.text)
                        for ch in list(node):
                            tag = ch.tag
                            if isinstance(tag, str) and tag.endswith('line-break'):
                                parts.append('\n')
                            elif isinstance(tag, str) and tag.endswith('tab'):
                                parts.append('\t')
                            else:
                                parts.append(extract_text(ch))
                            if ch.tail:
                                parts.append(ch.tail)
                    except Exception:
                        pass
                    return ''.join(parts)

                html_parts: list[str] = []
                for el in list(office_text):
                    t = el.tag
                    local = t.split('}')[-1] if isinstance(t, str) else ''
                    if local == 'h':
                        lvl = el.get(f"{{{ns['text']}}}outline-level", '1')
                        try:
                            lvl_i = int(lvl)
                        except Exception:
                            lvl_i = 1
                        lvl_i = max(1, min(6, lvl_i))
                        text_content = extract_text(el)
                        inner = escape(text_content).replace('\n', '<br/>')
                        html_parts.append(f'<h{lvl_i}>' + inner + f'</h{lvl_i}>')
                    elif local == 'p':
                        text_content = extract_text(el)
                        inner = escape(text_content).replace('\n', '<br/>')
                        html_parts.append('<p>' + inner + '</p>')
                    else:
                        # Skip unsupported element types for now
                        pass
                body_html = '\n'.join(html_parts) if html_parts else '<p></p>'
                # Use simple HTML structure for QTextEdit compatibility
                self.editor.blockSignals(True)
                self.set_editor_text(body_html)
                self.editor.blockSignals(False)
                self.is_rich_mode = True
                print(f"Rendered ODT via legacy parser; HTML length={len(body_html)}")
                try:
                    self.doc.set_text(self.editor.toPlainText())
                except Exception as e:
                    print('Backend set_text error:', e)
                return
            except Exception as e:
                print('ODT render failed:', e)
                QMessageBox.critical(self, "Open Error", f"Failed to parse ODT file: {e}")
            return
        try:
            self.doc.open(path)
            self.is_rich_mode = False
            self.set_editor_text(self.doc.get_text())
        except Exception as e:
            QMessageBox.critical(self, "Open Error", str(e))

    def save_file_dialog(self):
        path, _ = QFileDialog.getSaveFileName(self, "Save Document", os.getcwd(), "Documents (*.txt *.docx *.odt)")
        if not path:
            return
        try:
            self.doc.save(path)
        except Exception as e:
            QMessageBox.critical(self, "Save Error", str(e))

    def handle_command(self, cmd: Command):
        print(f"handle_command: received command {cmd.name}")
        # Ensure this runs on the main GUI thread
        if threading.current_thread() is not threading.main_thread():
            # This should not happen with QTimer.singleShot, but as a safeguard:
            print("handle_command: not on main thread, re-queuing")
            self.bridge.command_received.emit(cmd)
            return

        print(f"handle_command: processing {cmd.name} on main thread with args: {cmd.args}")
        name = cmd.name
        args = cmd.args or {}
        try:
            if name == "set_text":
                text = str(args.get("text", ""))
                print(f"handle_command: setting text to: {text[:50]}...")
                self.doc.set_text(text)
                # Reset rich mode when setting text via CLI to allow formatting
                self.is_rich_mode = False
                self.set_editor_text(text)
                print("handle_command: set_editor_text called")
            elif name == "insert_text":
                offset = int(args.get("offset", 0))
                text = str(args.get("text", ""))
                self.doc.insert_text(offset, text)
                # Reset rich mode when inserting text via CLI to allow formatting
                self.is_rich_mode = False
                self.set_editor_text(self.doc.get_text())
            elif name == "open":
                path = str(args.get("path", ""))
                if not path:
                    raise ValueError("path is required")
                self.open_path(path)
            elif name == "save":
                path = str(args.get("path", ""))
                if not path:
                    raise ValueError("path is required")
                self.doc.save(path)
            else:
                print(f"Unknown command: {name}")
        except Exception as e:
            print(f"handle_command: Error processing command {name}: {e}")
            QMessageBox.critical(self, "Command Error", f"{name}: {e}")

    def _maybe_show_bubble(self):
        cursor = self.editor.textCursor()
        has_selection = cursor.hasSelection()
        self.format_bubble.setVisible(has_selection)
        if not has_selection:
            return
        rect = self.editor.cursorRect(cursor)
        top_left = self.editor.viewport().mapTo(self.centralWidget(), rect.topLeft())
        x = top_left.x() - self.format_bubble.width() // 2 + rect.width() // 2
        y = max(0, top_left.y() - self.format_bubble.height() - 8)
        self.format_bubble.move(x, y)
        self.format_bubble.raise_()
        self._sync_bubble_state()

    def _sync_bubble_state(self):
        """Update check states of inline bubble actions to match current selection."""
        try:
            fmt = self.editor.currentCharFormat()
            self.act_bold.setChecked(fmt.fontWeight() > 50)
            self.act_italic.setChecked(fmt.fontItalic())
            self.act_underline.setChecked(fmt.fontUnderline())
        except Exception:
            pass
        try:
            al = self.editor.alignment()
            self.act_align_left.setChecked(bool(al & Qt.AlignLeft))
            self.act_align_center.setChecked(bool(al & Qt.AlignHCenter))
            self.act_align_right.setChecked(bool(al & Qt.AlignRight))
            self.act_align_justify.setChecked(bool(al & Qt.AlignJustify))
        except Exception:
            pass

    def eventFilter(self, obj, event):
        if obj is self.editor.viewport():
            if event.type() in (QEvent.MouseButtonPress, QEvent.Wheel, QEvent.KeyPress):
                QTimer.singleShot(0, lambda: self._maybe_show_bubble())
        return super().eventFilter(obj, event)

    def closeEvent(self, event: QCloseEvent) -> None:
        try:
            # Unregister from shared server if we used it
            if hasattr(self, 'assigned_port') and self.assigned_port is not None:
                from shared_server import get_shared_server
                shared_server = get_shared_server()
                shared_server.unregister_app("words")
                print("Words unregistered from shared server")
            # Stop fallback server if we used it
            elif hasattr(self, 'server'):
                self.server.stop()
        except Exception as e:
            print(f"Error during cleanup: {e}")
        finally:
            super().closeEvent(event)


def main():
    app = QApplication(sys.argv)
    win = MainWindow()
    win.resize(900, 700)
    win.show()
    # If a file path is provided as an argument, open it after the UI is shown
    try:
        args = list(sys.argv[1:])
        for a in args:
            try:
                if a and Path(a).exists():
                    QTimer.singleShot(0, lambda p=a: win.open_path(p))
                    break
            except Exception:
                continue
    except Exception:
        pass
    sys.exit(app.exec())


if __name__ == "__main__":
    main()